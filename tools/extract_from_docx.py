import re
import sys
from pathlib import Path

try:
    from docx import Document  # python-docx
except Exception as exc:
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    raise


FENCE_PATTERN = re.compile(r"^```(\w+)?\s*$")


def sanitize_filename(name: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "-", name.strip())
    return safe[:80] or "snippet"


def guess_extension(lang_hint: str) -> str:
    if not lang_hint:
        return ".txt"
    lang = lang_hint.lower()
    return {
        "c": ".c",
        "cpp": ".cpp",
        "h": ".h",
        "hpp": ".hpp",
        "ino": ".ino",
        "arduino": ".ino",
        "c++": ".cpp",
        "python": ".py",
        "py": ".py",
        "txt": ".txt",
        "js": ".js",
        "ts": ".ts",
        "json": ".json",
        "yaml": ".yaml",
        "yml": ".yml",
        "ini": ".ini",
        "env": ".env",
    }.get(lang, ".txt")


def extract_fenced_blocks(lines):
    blocks = []
    in_block = False
    lang = None
    buf = []
    for line in lines:
        if not in_block:
            m = FENCE_PATTERN.match(line)
            if m:
                in_block = True
                lang = m.group(1)
                buf = []
            continue
        else:
            if FENCE_PATTERN.match(line):
                blocks.append((lang, "".join(buf)))
                in_block = False
                lang = None
                buf = []
            else:
                buf.append(line)
    return blocks


def is_monospace(paragraph) -> bool:
    for run in paragraph.runs:
        font = run.font
        if font and font.name and str(font.name).lower() in {"courier", "courier new", "consolas", "lucida console", "monospace"}:
            return True
    return False


def extract_docx(docx_path: Path, out_dir: Path) -> int:
    doc = Document(str(docx_path))

    # First pass: capture plain text to detect fenced blocks
    full_text_lines = []
    for p in doc.paragraphs:
        full_text_lines.append(p.text + "\n")

    fenced = extract_fenced_blocks(full_text_lines)

    # Second pass: heuristics for code-like paragraphs (monospace, code styles)
    mono_blocks = []
    current = []
    for p in doc.paragraphs:
        if is_monospace(p) or p.style and "code" in p.style.name.lower():
            current.append(p.text)
        else:
            if current:
                mono_blocks.append("\n".join(current) + "\n")
                current = []
    if current:
        mono_blocks.append("\n".join(current) + "\n")

    out_dir.mkdir(parents=True, exist_ok=True)

    count = 0

    # Write fenced blocks with extensions by language hint
    for idx, (lang, content) in enumerate(fenced, start=1):
        ext = guess_extension(lang)
        fname = out_dir / f"fenced_{idx:02d}{ext}"
        fname.write_text(content, encoding="utf-8")
        count += 1

    # Write monospace-derived blocks as txt
    for idx, content in enumerate(mono_blocks, start=1):
        fname = out_dir / f"mono_{idx:02d}.txt"
        fname.write_text(content, encoding="utf-8")
        count += 1

    return count


def main():
    if len(sys.argv) < 2:
        print("Usage: python tools/extract_from_docx.py <path-to-docx> [out_dir]")
        sys.exit(2)
    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Not found: {docx_path}")
        sys.exit(1)
    out_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("extracted")
    num = extract_docx(docx_path, out_dir)
    print(f"Extracted {num} block(s) to {out_dir}")


if __name__ == "__main__":
    main()


